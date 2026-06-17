"""
build_rag_embeddings.py — Build Chroma vector index for a knowledge base.

Usage:
    source langgraph_env/bin/activate
    python scripts/build_rag_embeddings.py ./ragbase/rnai_records

Validates folder structure (metadata.md, raw/, vector/), chunks raw files
according to the configured chunk_strategy, generates embeddings with
BAAI/bge-large-en-v1.5, and stores results in Chroma.
"""

from __future__ import annotations

import json
import logging
import sys
import time
from pathlib import Path

_PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_PROJECT_ROOT))

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("build_rag")

# ============================================================
# Paths
# ============================================================
from tool_config import RAG_CHROMA_DIR

CHROMA_DIR = RAG_CHROMA_DIR

# ============================================================
# Allowed file formats
# ============================================================
ALLOWED_EXTENSIONS = {".txt", ".csv", ".json", ".md", ".docx", ".xlsx", ".pdf"}

# ============================================================
# Default metadata values
# ============================================================
DEFAULT_VALUES = {
    "description": None,  # filled dynamically with folder name
    "chunk_strategy": "general",
    "chunk_params": {},
    "recommended_search_mode": "hybrid",
    "recommended_semantic_weight": 0.4,
}

# ============================================================
# Chunk structure
# ============================================================
class Chunk:
    """A single retrievable text chunk with metadata."""

    def __init__(self, content: str, source: str, section: str = "",
                 parent_content: str = ""):
        self.content = content.strip()
        self.source = source
        self.section = section
        self.parent_content = parent_content
        self.chunk_id = ""

    def to_dict(self) -> dict:
        return {
            "content": self.content,
            "source": self.source,
            "section": self.section,
            "parent_content": self.parent_content,
        }


# ============================================================
# File parsers
# ============================================================
def _read_docx(path: Path) -> list[str]:
    from docx import Document
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _read_xlsx_rows(path: Path) -> tuple[list[str], list[list[str]]]:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return [], []
    header = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
    data: list[list[str]] = []
    for row in rows[1:]:
        vals = [str(v) if v is not None else "" for v in row]
        if any(v.strip() for v in vals):
            data.append(vals)
    return header, data


def _read_pdf(path: Path) -> list[str]:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text.strip():
            pages.append(text.strip())
    return pages


# ============================================================
# Chunk strategy: general
# ============================================================
def _chunk_general(file_path: Path) -> list[Chunk]:
    suffix = file_path.suffix.lower()
    name = file_path.name
    chunks: list[Chunk] = []

    if suffix == ".xlsx":
        header, rows = _read_xlsx_rows(file_path)
        if not header:
            return chunks
        for row in rows:
            fields = [f"{h}: {v}" for h, v in zip(header, row) if v.strip()]
            content = " | ".join(fields)
            if content:
                chunks.append(Chunk(content, name, "record"))

    elif suffix == ".docx":
        paragraphs = _read_docx(file_path)
        current_lines: list[str] = []
        current_title = ""
        for para in paragraphs:
            if para == "{}":
                if current_lines:
                    chunks.append(Chunk("\n".join(current_lines), name, current_title))
                    current_lines = []
                current_title = ""
            elif current_title == "" and not para.startswith(("{", "$")):
                current_title = para
            else:
                current_lines.append(para)
        if current_lines:
            chunks.append(Chunk("\n".join(current_lines), name, current_title))

    elif suffix == ".pdf":
        pages = _read_pdf(file_path)
        for i, text in enumerate(pages):
            chunks.append(Chunk(text, name, f"page_{i + 1}"))

    elif suffix in (".txt", ".csv", ".json", ".md"):
        text = file_path.read_text(encoding="utf-8", errors="replace")
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        for i, para in enumerate(paragraphs):
            chunks.append(Chunk(para, name, f"paragraph_{i + 1}"))

    return chunks


# ============================================================
# Chunk strategy: parent_child
# ============================================================
def _chunk_parent_child(file_path: Path, params: dict) -> list[Chunk]:
    suffix = file_path.suffix.lower()
    name = file_path.name
    parent_marker = params.get("parent_marker", "\n\n")
    child_marker = params.get("child_marker", "\n")

    if suffix == ".docx":
        paragraphs = _read_docx(file_path)
        if parent_marker in ("{}", "$$") or child_marker in ("$", "$$"):
            return _chunk_docx_with_markers(file_path, parent_marker, child_marker)
        text = "\n".join(paragraphs)
        return _chunk_text_parent_child(text, name, parent_marker, child_marker)

    elif suffix in (".txt", ".csv", ".json", ".md"):
        text = file_path.read_text(encoding="utf-8", errors="replace")
        return _chunk_text_parent_child(text, name, parent_marker, child_marker)

    elif suffix == ".xlsx":
        header, rows = _read_xlsx_rows(file_path)
        if not header:
            return chunks
        for row in rows:
            if not row[0].strip():
                continue
            parent_text = row[0].strip()
            child_text = " | ".join(
                f"{h}: {v}" for h, v in zip(header[1:], row[1:]) if v.strip()
            )
            if child_text:
                chunks.append(Chunk(
                    content=child_text, source=name,
                    section=parent_text, parent_content=parent_text,
                ))
        return chunks

    return []


def _chunk_docx_with_markers(path: Path, parent_marker: str,
                              child_marker: str) -> list[Chunk]:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    chunks: list[Chunk] = []
    pending_new_parent = False
    pending_new_child = False
    parent_title = ""
    parent_lines: list[str] = []
    child_title = ""
    child_lines: list[str] = []

    def _flush_parent():
        nonlocal parent_lines
        if parent_lines:
            chunks.append(Chunk("\n".join(parent_lines), path.name, parent_title))
            parent_lines = []

    def _flush_child():
        nonlocal child_lines
        if child_lines and parent_title:
            content = "\n".join(child_lines)
            parent_content = "\n".join(parent_lines) if parent_lines else ""
            chunks.append(Chunk(content, path.name,
                                f"{parent_title} / {child_title}", parent_content))
            child_lines = []

    for para in paragraphs:
        if para == parent_marker:
            _flush_child()
            _flush_parent()
            pending_new_parent = True
            continue
        if para == child_marker:
            _flush_child()
            pending_new_child = True
            continue
        if pending_new_parent:
            parent_title = para
            pending_new_parent = False
            pending_new_child = False
            continue
        if pending_new_child:
            child_title = para
            pending_new_child = False
            continue
        if child_title:
            child_lines.append(para)
        elif parent_title:
            parent_lines.append(para)
        else:
            parent_lines.append(para)
    _flush_child()
    _flush_parent()
    return chunks


def _chunk_text_parent_child(text: str, source: str,
                              parent_marker: str, child_marker: str) -> list[Chunk]:
    chunks: list[Chunk] = []
    parents = text.split(parent_marker)
    for parent_block in parents:
        if not parent_block.strip():
            continue
        lines = parent_block.strip().split("\n")
        parent_title = lines[0] if lines else ""
        parent_body = "\n".join(lines[1:]) if len(lines) > 1 else ""
        if child_marker == "\n":
            for line in lines[1:]:
                if line.strip():
                    chunks.append(Chunk(
                        content=line.strip(), source=source,
                        section=f"{parent_title} / {line.strip()[:50]}",
                        parent_content=parent_block.strip(),
                    ))
        else:
            children = parent_body.split(child_marker)
            for child_block in children:
                if child_block.strip():
                    child_lines = child_block.strip().split("\n")
                    child_title = child_lines[0] if child_lines else ""
                    child_body = "\n".join(child_lines[1:]) if len(child_lines) > 1 else ""
                    chunks.append(Chunk(
                        content=child_body if child_body else child_title,
                        source=source,
                        section=f"{parent_title} / {child_title}",
                        parent_content=parent_block.strip(),
                    ))
    return chunks


# ============================================================
# Chunk strategy: qa
# ============================================================
def _chunk_qa(file_path: Path, params: dict) -> list[Chunk]:
    suffix = file_path.suffix.lower()
    name = file_path.name
    question_col = params.get("question_col", "Question")
    answer_col = params.get("answer_col", "Answer")
    chunks: list[Chunk] = []

    if suffix == ".xlsx":
        header, rows = _read_xlsx_rows(file_path)
        if not header:
            return chunks
        try:
            q_idx = header.index(question_col)
            a_idx = header.index(answer_col)
        except ValueError:
            logger.warning("QA columns not found in %s (need '%s' and '%s')",
                           name, question_col, answer_col)
            return chunks
        for row in rows:
            q_text = row[q_idx].strip() if q_idx < len(row) else ""
            a_text = row[a_idx].strip() if a_idx < len(row) else ""
            if q_text and a_text:
                chunks.append(Chunk(content=a_text, source=name, section=q_text))

    elif suffix == ".csv":
        import csv, io
        reader = csv.DictReader(io.StringIO(
            file_path.read_text(encoding="utf-8", errors="replace")))
        for row in reader:
            q_text = row.get(question_col, "").strip()
            a_text = row.get(answer_col, "").strip()
            if q_text and a_text:
                chunks.append(Chunk(content=a_text, source=name, section=q_text))

    elif suffix in (".docx", ".txt", ".json", ".md"):
        paragraphs = (_read_docx(file_path) if suffix == ".docx"
                      else [p.strip() for p in
                            file_path.read_text(encoding="utf-8").split("\n")
                            if p.strip()])
        current_q = ""
        current_a_lines: list[str] = []
        for para in paragraphs:
            if para.startswith("Q:") or para.startswith("Question:"):
                if current_q and current_a_lines:
                    chunks.append(Chunk(
                        content="\n".join(current_a_lines),
                        source=name, section=current_q,
                    ))
                current_q = para.split(":", 1)[1].strip() if ":" in para else para
                current_a_lines = []
            elif para.startswith("A:") or para.startswith("Answer:"):
                a_text = para.split(":", 1)[1].strip() if ":" in para else para
                if a_text:
                    current_a_lines.append(a_text)
            elif current_q:
                current_a_lines.append(para)
        if current_q and current_a_lines:
            chunks.append(Chunk(
                content="\n".join(current_a_lines), source=name, section=current_q,
            ))

    return chunks


# ============================================================
# Chunk dispatch
# ============================================================
CHUNK_DISPATCH = {
    "general": _chunk_general,
    "parent_child": _chunk_parent_child,
    "qa": _chunk_qa,
}

# ============================================================
# Validation
# ============================================================
def validate_kb_folder(kb_path: Path) -> dict:
    """Validate the knowledge base folder structure and metadata.

    Returns parsed metadata dict on success.
    Raises SystemExit on validation failure.
    """
    kb_path = kb_path.resolve()

    if not kb_path.exists():
        logger.error("Folder not found: %s", kb_path)
        sys.exit(1)
    if not kb_path.is_dir():
        logger.error("Not a directory: %s", kb_path)
        sys.exit(1)

    folder_name = kb_path.name

    # ---- Check required files/folders ----
    md_path = kb_path / "metadata.md"
    if not md_path.exists():
        logger.error("metadata.md not found in %s", kb_path)
        sys.exit(1)

    raw_dir = kb_path / "raw"
    if not raw_dir.exists():
        logger.error("raw/ directory not found in %s", kb_path)
        sys.exit(1)
    if not raw_dir.is_dir():
        logger.error("raw is not a directory: %s", raw_dir)
        sys.exit(1)

    vector_dir = kb_path / "vector"
    if not vector_dir.exists():
        logger.error("vector/ directory not found in %s", kb_path)
        sys.exit(1)
    if not vector_dir.is_dir():
        logger.error("vector is not a directory: %s", vector_dir)
        sys.exit(1)

    # ---- Parse metadata.md ----
    try:
        import yaml
        raw = md_path.read_text(encoding="utf-8")
        if raw.startswith("---"):
            parts = raw.split("---", 2)
            if len(parts) >= 3:
                raw = parts[1]
        meta = yaml.safe_load(raw)
    except Exception as e:
        logger.error("Failed to parse metadata.md: %s", e)
        sys.exit(1)

    if not isinstance(meta, dict):
        logger.error("metadata.md is not a valid YAML dictionary")
        sys.exit(1)

    # ---- Validate name field ----
    name = meta.get("name")
    if name is None or (isinstance(name, str) and not name.strip()):
        logger.error("metadata.md 'name' field is empty — cannot build KB")
        sys.exit(1)

    if name != folder_name:
        logger.warning(
            "metadata.md name='%s' differs from folder name='%s'. "
            "Using metadata.md name '%s' for the Chroma collection.",
            name, folder_name, name,
        )

    # ---- Fill defaults for other fields ----
    defaults = dict(DEFAULT_VALUES)
    defaults["description"] = (
        f'A useful knowledge base when you need to search for content about "{name}".'
    )

    for field, default_val in defaults.items():
        if field not in meta or meta[field] is None:
            meta[field] = default_val
            logger.info("Field '%s' using default: %s", field, default_val)
        elif field == "description" and isinstance(meta[field], str) and not meta[field].strip():
            meta[field] = default_val
            logger.info("Field 'description' is empty, using default")

    # ---- Validate chunk_strategy ----
    valid_strategies = {"general", "parent_child", "qa"}
    if meta["chunk_strategy"] not in valid_strategies:
        logger.error(
            "Unsupported chunk_strategy: '%s'. Valid options: %s",
            meta["chunk_strategy"], ", ".join(sorted(valid_strategies)),
        )
        sys.exit(1)

    # ---- Check raw/ files ----
    raw_files = sorted([
        f for f in raw_dir.iterdir()
        if f.is_file() and not f.name.startswith(".")
    ])

    if not raw_files:
        logger.error("raw/ directory is empty — cannot build KB")
        sys.exit(1)

    # Check for unsupported formats
    unsupported = [f for f in raw_files if f.suffix.lower() not in ALLOWED_EXTENSIONS]
    if unsupported:
        bad_exts = sorted(set(f.suffix for f in unsupported))
        logger.error(
            "Unsupported file formats in raw/. Allowed: %s\n"
            "  Found: %s\n"
            "  Files: %s",
            ", ".join(sorted(ALLOWED_EXTENSIONS)),
            ", ".join(bad_exts),
            ", ".join(f.name for f in unsupported),
        )
        sys.exit(1)

    # ---- Check vector/ files ----
    vector_files = sorted([
        f for f in vector_dir.iterdir()
        if f.is_file() and not f.name.startswith(".")
    ])
    if vector_files:
        print(f"⚠️  vector/ directory already contains files. Building will overwrite. Continue? (y/N): ", end="", flush=True)
        try:
            answer = input().strip().lower()
        except (EOFError, KeyboardInterrupt):
            answer = "n"
        if answer != "y":
            logger.info("User cancelled build")
            sys.exit(0)

    return meta, name, kb_path, raw_dir, vector_dir


# ============================================================
# Build index
# ============================================================
def build_index(kb_path_str: str):
    """Validate and build Chroma index for a single knowledge base folder."""

    kb_path = Path(kb_path_str)
    meta, coll_name, kb_dir, raw_dir, vector_dir = validate_kb_folder(kb_path)

    strategy = meta["chunk_strategy"]
    chunker = CHUNK_DISPATCH[strategy]
    chunk_params = meta.get("chunk_params", {})

    logger.info("=== Building KB '%s' (strategy: %s) ===", coll_name, strategy)

    # ---- Step 1: Chunk ----
    logger.info("Step 1: Chunking files in raw/ ...")
    all_chunks: list[Chunk] = []
    raw_files = sorted([
        f for f in raw_dir.iterdir()
        if f.is_file() and not f.name.startswith(".")
    ])

    for fpath in raw_files:
        try:
            if strategy == "general":
                new_chunks = chunker(fpath)
            else:
                new_chunks = chunker(fpath, chunk_params)
            all_chunks.extend(new_chunks)
            logger.info("  %s → %d chunks", fpath.name, len(new_chunks))
        except Exception:
            logger.exception("Error chunking %s", fpath)
            continue

    if not all_chunks:
        logger.error("No chunks generated — cannot build KB")
        sys.exit(1)

    logger.info("Total: %d chunks", len(all_chunks))

    # Write chunk JSON backup to vector/
    vector_dir.mkdir(parents=True, exist_ok=True)
    chunks_data = [c.to_dict() for c in all_chunks]
    backup_path = vector_dir / "chunks.json"
    backup_path.write_text(
        json.dumps(chunks_data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    logger.info("Chunk backup written to: %s", backup_path)

    # ---- Step 2: Embed ----
    logger.info("Step 2: Loading embedding model (BAAI/bge-large-en-v1.5)...")
    from sentence_transformers import SentenceTransformer
    t0 = time.time()
    model = SentenceTransformer("BAAI/bge-large-en-v1.5")
    logger.info("Model loaded (%.1fs)", time.time() - t0)

    texts = [c.content for c in all_chunks]
    logger.info("Generating %d embeddings...", len(texts))
    t0 = time.time()
    embeddings = model.encode(texts, show_progress_bar=True, batch_size=32)
    logger.info("Embeddings done (%.1fs, shape: %s)", time.time() - t0, embeddings.shape)

    # ---- Step 3: Store in Chroma ----
    logger.info("Step 3: Writing to Chroma DB...")
    import chromadb
    from chromadb.config import Settings as ChromaSettings

    chroma_client = chromadb.PersistentClient(
        path=str(CHROMA_DIR),
        settings=ChromaSettings(anonymized_telemetry=False),
    )

    # Delete existing collection if exists
    try:
        chroma_client.delete_collection(coll_name)
    except Exception:
        pass

    collection = chroma_client.create_collection(
        name=coll_name,
        metadata={
            "hnsw:space": "cosine",
            "description": meta.get("description", ""),
            "chunk_strategy": strategy,
        },
    )

    ids: list[str] = []
    metadatas: list[dict] = []
    documents: list[str] = []

    for i, (chunk, emb) in enumerate(zip(all_chunks, embeddings)):
        cid = f"{coll_name}_{i:06d}"
        chunk.chunk_id = cid
        ids.append(cid)
        metadatas.append({
            "source": chunk.source,
            "section": chunk.section,
            "has_parent": bool(chunk.parent_content),
            "parent_content": chunk.parent_content if chunk.parent_content else "",
        })
        documents.append(chunk.content)

    collection.add(
        ids=ids,
        embeddings=[emb.tolist() for emb in embeddings],
        metadatas=metadatas,
        documents=documents,
    )

    logger.info("=== Build complete: '%s' (%d chunks written to %s) ===",
                coll_name, len(all_chunks), CHROMA_DIR)


# ============================================================
# CLI
# ============================================================
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/build_rag_embeddings.py <kb_folder_path>")
        print("Example: python scripts/build_rag_embeddings.py ./ragbase/your_kb_folder")
        sys.exit(1)

    build_index(sys.argv[1])
